from langchain.schema import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import os
from typing import List, Optional, Dict, Any
from utils.logging_config import logger
import tempfile

class DocumentValidator:
    """Handles document validation and preprocessing."""
    
    def __init__(self):
        self.logger = logger
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
    
    def validate_file_path(self, file_path: str) -> bool:
        """Validate file exists and has supported extension."""
        if not os.path.exists(file_path):
            self.logger.error(f"File not found: {file_path}")
            return False
            
        extension = os.path.splitext(file_path)[1].lower()
        if extension not in self.supported_extensions:
            self.logger.error(f"Unsupported file type: {extension}")
            return False
            
        return True
    
    def validate_content(self, content: str) -> bool:
        """Validate document content."""
        if not content or not content.strip():
            self.logger.warning("Empty document content")
            return False
        return True

class PDFLoader:
    """Handles loading and processing PDF documents."""
    
    def __init__(self):
        self.logger = logger
    
    def load(self, file_path: str) -> Optional[str]:
        """Load and extract text from PDF."""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error loading PDF {file_path}: {e}")
            return None

class DocxLoader:
    """Handles loading and processing DOCX documents."""
    
    def __init__(self):
        self.logger = logger
    
    def load(self, file_path: str) -> Optional[str]:
        """Load and extract text from DOCX."""
        try:
            doc = DocxDocument(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
        except Exception as e:
            self.logger.error(f"Error loading DOCX {file_path}: {e}")
            return None

class TextLoader:
    """Handles loading text files."""
    
    def __init__(self):
        self.logger = logger
    
    def load(self, file_path: str) -> Optional[str]:
        """Load text file content."""
        try:
            # Verify file extension
            if not file_path.lower().endswith('.txt'):
                self.logger.error(f"Invalid file extension for text file: {file_path}")
                return None
                
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            if not content.strip():
                self.logger.warning(f"Empty content in text file: {file_path}")
                return None
                
            return content
                
        except Exception as e:
            self.logger.error(f"Error loading text file {file_path}: {e}")
            return None

class DocumentLoader:
    """Main document loading coordinator."""
    
    def __init__(self):
        self.validator = DocumentValidator()
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DocxLoader(),
            '.txt': TextLoader()
        }
        self.logger = logger
    
    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """Load multiple documents with validation."""
        documents = []
        
        for file_path in file_paths:
            doc = self.load_single_document(file_path)
            if doc:
                documents.append(doc)
                
        return documents
    
    def load_single_document(self, file_path: str) -> Optional[Document]:
        """Load a single document with validation."""
        try:
            # Get file extension
            extension = os.path.splitext(file_path)[1].lower()
            if not extension:
                self.logger.error(f"No file extension found for: {file_path}")
                return None
            
            # Get appropriate loader
            loader = self.loaders.get(extension)
            if not loader:
                self.logger.error(f"No loader found for extension: {extension}")
                return None
            
            # Load content
            content = loader.load(file_path)
            if not content or not self.validator.validate_content(content):
                return None
            
            # Create document
            return Document(
                page_content=content,
                metadata=self._create_metadata(file_path)
            )
            
        except Exception as e:
            self.logger.error(f"Error loading document {file_path}: {e}")
            return None
    
    def _create_metadata(self, file_path: str) -> Dict[str, Any]:
        """Create metadata for document."""
        return {
            "source": file_path,
            "file_type": os.path.splitext(file_path)[1].lower(),
            "file_name": os.path.basename(file_path)
        }

def load_documents(file_paths: List[str]) -> List[Document]:
    """Load documents from file paths."""
    loader = DocumentLoader()
    return loader.load_documents(file_paths)

def process_with_dspy(documents: List[Document]) -> List[Document]:
    """Process documents using the DSPy pipeline."""
    from dspy_pipeline import IEPPipeline
    
    pipeline = IEPPipeline()
    results = []
    
    for doc in documents:
        try:
            processed_docs = pipeline.process_documents([doc])
            
            if processed_docs and len(processed_docs) > 0:
                results.append(processed_docs[0])
                logger.debug(f"Successfully processed document: {doc.metadata.get('source')}")
            else:
                logger.warning(f"No processing results for document: {doc.metadata.get('source')}")
                results.append(doc)
            
        except Exception as e:
            logger.error(f"Error processing document with DSPy: {e}")
            results.append(doc)
    
    return results

class FileUploadHandler:
    """Handles file upload processing and temporary file management."""
    
    def __init__(self):
        self.logger = logger
        self.temp_files = []
    
    def process_uploaded_file(self, uploaded_file) -> Optional[str]:
        """Process an uploaded file and return its temporary path."""
        try:
            import tempfile
            
            # Get file extension from original name
            extension = os.path.splitext(uploaded_file.name)[1]
            if not extension:
                extension = '.txt'  # Default to .txt if no extension
            
            # Create temporary file with extension
            temp_file = tempfile.NamedTemporaryFile(
                suffix=extension,
                delete=False
            )
            self.temp_files.append(temp_file.name)
            
            # Write uploaded file content to temporary file
            temp_file.write(uploaded_file.getvalue())
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            self.logger.error(f"Error processing uploaded file: {e}")
            return None
    
    def cleanup(self):
        """Clean up temporary files."""
        for temp_path in self.temp_files:
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except Exception as e:
                self.logger.error(f"Error cleaning up temp file {temp_path}: {e}")